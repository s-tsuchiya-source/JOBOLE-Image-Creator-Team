from __future__ import annotations

import csv
import json
from pathlib import Path

import yaml
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}
DOCUMENT_EXTENSIONS = {".docx", ".xlsx", ".pdf"}
REFERENCE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".webp", ".gif", ".svg", ".pdf", ".pptx"
}
SUPPORTED_SOURCE_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"文字コードを判定できません: {path}")


def extract_csv(path: Path) -> str:
    text = read_text_file(path)
    return "\n".join("\t".join(row) for row in csv.reader(text.splitlines()))


def extract_json(path: Path) -> str:
    data = json.loads(read_text_file(path))
    return json.dumps(data, ensure_ascii=False, indent=2)


def extract_yaml(path: Path) -> str:
    data = yaml.safe_load(read_text_file(path))
    return yaml.safe_dump(data, allow_unicode=True, sort_keys=False)


def extract_docx(path: Path) -> str:
    document = Document(path)
    chunks = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            chunks.append(value)
    for table in document.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(chunks)


def extract_xlsx(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    chunks = []
    try:
        for worksheet in workbook.worksheets:
            chunks.append(f"## Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    chunks.append("\t".join(values))
    finally:
        workbook.close()
    return "\n".join(chunks)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(f"## Page {index}\n{text.strip()}")
    if not chunks:
        raise ValueError(
            "PDFからテキストを抽出できません。画像PDF/スキャンPDFの可能性があります。"
        )
    return "\n\n".join(chunks)


def extract_source_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return read_text_file(path)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix == ".json":
        return extract_json(path)
    if suffix in {".yaml", ".yml"}:
        return extract_yaml(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"未対応形式です: {path.suffix}")


def list_files(folder: Path) -> list[Path]:
    if not folder.exists():
        return []
    return sorted(path for path in folder.rglob("*") if path.is_file())


def normalize_project_inputs(project_dir: Path) -> dict:
    request_dir = project_dir / "00_request"
    job_dir = request_dir / "inbox" / "job_posting"
    hearing_dir = request_dir / "inbox" / "hearing"
    reference_dir = request_dir / "inbox" / "references"
    normalized_dir = request_dir / "normalized"

    for folder in (job_dir, hearing_dir, reference_dir, normalized_dir):
        folder.mkdir(parents=True, exist_ok=True)

    categories = {
        "job_posting": list_files(job_dir),
        "hearing": list_files(hearing_dir),
        "references": list_files(reference_dir),
    }

    index = {
        "job_posting": [],
        "hearing": [],
        "references": [],
        "errors": [],
    }
    bundle_parts = ["# Source Bundle", ""]

    for category in ("job_posting", "hearing"):
        for path in categories[category]:
            relative = path.relative_to(project_dir).as_posix()
            item = {
                "path": relative,
                "name": path.name,
                "extension": path.suffix.lower(),
            }
            if path.suffix.lower() not in SUPPORTED_SOURCE_EXTENSIONS:
                item["status"] = "unsupported"
                index[category].append(item)
                index["errors"].append(f"未対応形式: {relative}")
                continue
            try:
                text = extract_source_text(path).strip()
                item["status"] = "loaded"
                item["characters"] = len(text)
                index[category].append(item)
                bundle_parts.extend(
                    [
                        f"# {category}: {path.name}",
                        f"source_path: {relative}",
                        "",
                        text,
                        "",
                        "---",
                        "",
                    ]
                )
            except Exception as exc:
                item["status"] = "error"
                item["error"] = str(exc)
                index[category].append(item)
                index["errors"].append(f"{relative}: {exc}")

    for path in categories["references"]:
        relative = path.relative_to(project_dir).as_posix()
        index["references"].append(
            {
                "path": relative,
                "name": path.name,
                "extension": path.suffix.lower(),
                "status": "indexed",
            }
        )

    bundle_path = normalized_dir / "source-bundle.md"
    index_path = normalized_dir / "source-index.json"

    bundle_path.write_text("\n".join(bundle_parts), encoding="utf-8-sig")
    index_path.write_text(
        json.dumps(index, ensure_ascii=False, indent=2),
        encoding="utf-8-sig",
    )

    loaded_jobs = sum(item.get("status") == "loaded" for item in index["job_posting"])
    loaded_hearing = sum(item.get("status") == "loaded" for item in index["hearing"])

    # A job posting is the only mandatory source. Hearing information is optional;
    # specialist agents must request clarification only when missing information
    # materially blocks production.
    ready = loaded_jobs > 0 and not index["errors"]

    return {
        "source_bundle": str(bundle_path),
        "source_index": str(index_path),
        "job_posting_count": loaded_jobs,
        "hearing_count": loaded_hearing,
        "hearing_provided": loaded_hearing > 0,
        "reference_count": len(index["references"]),
        "errors": index["errors"],
        "ready": ready,
    }
